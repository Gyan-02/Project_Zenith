from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("submission")
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "Project_Zenith_Round2_README.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "E8EEF5"
GRAY_FILL = "F2F4F7"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(20, 20, 20)
    return p


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Project Zenith - Round 2 README")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Project Zenith - Round 2 README")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(
        "AI-powered sky digital twin with CesiumJS visualization, live celestial data, "
        "satellite pass predictions, and a Gemini-grounded cosmic narrator."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.6, 4.9])
    rows = [
        ("Project", "Project Zenith"),
        ("Stack", "Next.js, Node.js, Express, CesiumJS, LangChain, ChromaDB, Gemini 2.5 Flash"),
        ("Mode support", "Live provider mode plus demo fixture fallback for reliable presentations"),
    ]
    for row, (key, value) in zip(table.rows, rows):
        set_cell_fill(row.cells[0], GRAY_FILL)
        row.cells[0].paragraphs[0].add_run(key).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_dependency_table(doc):
    doc.add_heading("Dependencies", level=1)
    add_para(
        doc,
        "The project uses a workspace structure with separate frontend and backend packages. "
        "Node.js 20 or later is recommended for local development and deployment.",
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1.4, 2.0, 3.1])
    headers = ["Area", "Dependency", "Purpose"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_FILL)
        cell.paragraphs[0].add_run(header).bold = True

    rows = [
        ("Runtime", "Node.js 20+", "Runs the Next.js frontend and Express backend."),
        ("Frontend", "Next.js, React", "Application shell, panels, routing, and client UI."),
        ("Visualization", "CesiumJS", "Interactive sky/globe rendering and object/entity display."),
        ("Backend", "Express, Zod", "HTTP APIs, validation, and typed request/response contracts."),
        ("AI", "Gemini 2.5 Flash, LangChain", "Grounded cosmic narrator and intent handling."),
        ("Knowledge", "ChromaDB", "Semantic retrieval for astronomy/cultural reference context."),
        ("Orbital data", "satellite.js, CelesTrak", "TLE propagation and satellite/ISS pass prediction."),
        ("Weather", "OpenWeather API", "Live observing conditions such as clouds, humidity, wind, and temperature."),
    ]
    for area, dep, purpose in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = dep
        cells[2].text = purpose
        for cell in cells:
            set_cell_margins(cell)


def build_doc():
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("Installation and Setup Instructions", level=1)
    add_para(doc, "1. Clone the repository and enter the project folder:")
    add_code_block(doc, ["git clone <your-repository-url>", "cd zenith"])
    add_para(doc, "2. Install workspace dependencies from the repository root:")
    add_code_block(doc, ["npm.cmd install"])
    add_para(doc, "3. Create environment files:")
    add_code_block(doc, ["Copy-Item backend\\.env.example backend\\.env", "Copy-Item frontend\\.env.example frontend\\.env.local"])
    add_para(doc, "4. Add required and optional environment variables.")

    env_table = doc.add_table(rows=1, cols=3)
    env_table.style = "Table Grid"
    set_table_width(env_table, [2.4, 1.1, 3.0])
    for idx, header in enumerate(["Variable", "Where", "Notes"]):
        cell = env_table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT_FILL)
        cell.paragraphs[0].add_run(header).bold = True
    env_rows = [
        ("NEXT_PUBLIC_API_URL", "frontend", "Backend URL, for example http://localhost:4000 locally or the deployed Render URL."),
        ("NEXT_PUBLIC_CESIUM_ION_TOKEN", "frontend", "Optional Cesium ion token for hosted terrain/imagery."),
        ("GEMINI_API_KEY", "backend", "Required for live Gemini narrator responses."),
        ("OPENWEATHER_API_KEY", "backend", "Optional live observing conditions provider."),
        ("CHROMA_API_KEY / CHROMA_URL", "backend", "Optional Chroma Cloud or local Chroma connection."),
        ("FRONTEND_ORIGIN", "backend", "Allowed CORS origin for the deployed frontend."),
    ]
    for row in env_rows:
        cells = env_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])

    add_para(doc, "5. Seed the knowledge base if using Chroma:")
    add_code_block(doc, ["npm.cmd run seed -w backend"])
    add_para(doc, "6. Run the application locally:")
    add_code_block(doc, ["npm.cmd run dev"])
    add_para(doc, "7. Open the app:")
    add_code_block(doc, ["Frontend: http://localhost:3000", "Backend health: http://localhost:4000/health"])
    add_para(doc, "8. Optional demo mode:")
    add_code_block(doc, ["http://localhost:3000/?demo=1"])

    doc.add_heading("Website Functionality and Unique Features", level=1)
    add_para(
        doc,
        "Project Zenith turns the current sky above an observer into an interactive digital twin. "
        "The user can explore celestial objects, ask natural-language questions, view pass predictions, "
        "and switch between live data and reliable demo fixtures.",
    )
    add_bullets(doc, [
        "Interactive CesiumJS sky visualization with planets, Moon, stars, ISS, satellites, paths, and labels.",
        "Mission-control UI with top command bar, left tool rail, bottom context tray, and compact narrator pill.",
        "Live sky-state API that combines observer location, current time, planetary data, bright stars, ISS, and satellites.",
        "Search and navigation system for finding objects such as Saturn, ISS, Vega, Moon, and visible satellites.",
        "Gemini-grounded Cosmic Narrator that answers astronomy questions using sky-state and knowledge retrieval context.",
        "Satellite pass predictions powered by CelesTrak TLE data and satellite.js orbital propagation.",
        "Sky events panel for meteor showers, eclipses, visibility windows, and conjunctions.",
        "ISS live video panel available as a deliberate one-click demo moment.",
        "Observing conditions panel using weather data to describe cloud cover, visibility, humidity, temperature, and wind.",
        "Demo mode fallback so the project remains presentable even if external APIs are slow or unavailable.",
        "Share and snapshot support for capturing a selected location, time, layers, and object context.",
        "Data provenance labels that explain whether information is live, demo, static, cached, or fallback.",
    ])

    doc.add_heading("What Makes Project Zenith Stand Out", level=2)
    add_bullets(doc, [
        "It combines real-time orbital mechanics, AI narration, and interactive 3D visualization in one product.",
        "The narrator is grounded by the current sky-state instead of giving generic astronomy answers.",
        "The UI is designed for a hackathon demo: major tools are one click away and do not require scrolling through a dashboard.",
        "Live/demo separation makes the product resilient during presentations while still supporting real provider data.",
    ])

    add_dependency_table(doc)

    doc.add_heading("Project Structure", level=1)
    add_code_block(doc, [
        "backend/   Node.js + Express APIs, Gemini narration, Chroma retrieval, sky-state services",
        "frontend/  Next.js UI, Cesium scene, panels, narrator, search, and demo/live controls",
        "data/      Demo fixtures and astronomy/cultural datasets",
        "docs/      Runbooks, demo scripts, smoke checklists, and project notes",
        "scripts/   Validation, doctor, and helper scripts",
    ])

    doc.add_heading("Deployment Notes", level=1)
    add_numbered(doc, [
        "Deploy the backend as a Node web service on Render using build command npm install && npm run build -w backend.",
        "Use start command npm run start -w backend and health check path /health.",
        "Deploy the frontend on Vercel or Render with NEXT_PUBLIC_API_URL pointing to the Render backend URL.",
        "Set FRONTEND_ORIGIN on the backend to the exact deployed frontend URL to satisfy CORS.",
        "After changing frontend environment variables, redeploy the frontend because NEXT_PUBLIC values are built into the client bundle.",
    ])

    doc.add_heading("Quick Verification Checklist", level=1)
    add_bullets(doc, [
        "Frontend loads without a blank screen.",
        "Globe and sky objects render.",
        "Search finds Saturn and ISS.",
        "Sky, Events, Passes, ISS, Layers, Data, and Save panels open from the rail.",
        "Narrator opens from the Ask Zenith pill and answers at least one question.",
        "Backend /health returns status ok.",
        "Demo mode works at /?demo=1.",
    ])

    doc.add_heading("Conclusion", level=1)
    add_para(
        doc,
        "Project Zenith is a complete AI-powered astronomy experience: a live sky digital twin, an orbital-data backend, "
        "and a grounded narrator wrapped in a polished mission-control interface. The system is designed to work with live APIs "
        "for realism while retaining demo fixtures for reliable judging and presentation.",
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
